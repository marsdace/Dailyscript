#!/usr/bin/env python
"""Django's command-line utility for administrative tasks."""
import os
import sys

from docx import Document
from docx.shared import Inches
import pandas as pd


def check_and_change(document, replace_dict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
   （key 和 value 都是replace_dict中的键值对。）
    """
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    print(key+"->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document

# 读取excel文件
df = pd.read_excel(".\data\一线一案材料.xlsx",sheet_name='Sheet1')
line = 0
while line< len(df) :
    # print(df.iloc[line])
    print(line)
    replace_dict = {
    # 线路名称	变电站	运维单位	公变总数	专变总数	11.8合格率	本月合格天数	不合格原因	治理措施		11.13日线损	是否治理完成
        "线路名称":str(df.iloc[line,1]),
        "变电站":str(df.iloc[line,2]),
        "公变总数": str(df.iloc[line,4]),
        "专变总数": str(df.iloc[line,5]),
        "应有公变数": str(df.iloc[line,6]),
        "应有专变数": str(df.iloc[line,7]),
        "11.8合格率": str(df.iloc[line,8]),
        "#本月合格天数#": str(df.iloc[line,9]),
        "#不合格原因#": str(df.iloc[line,10]),
        "治理措施": str(df.iloc[line,11]),
        "11.13日线损":str( df.iloc[line,12]),
        "是否治理完成": str(df.iloc[line,13]),
    }
    print(replace_dict)
    output_filename=df.iloc[line,1]+"路线路线损提质一线一案.docx"
    document = Document(".\data\线路线损提质一线一案.docx")  # 首先这是包的主要接口，这应该是利用的设计模式的一种，用来创建docx文档，里面也可以包含文档路径(d:\\2.docx)
    document = check_and_change(document, replace_dict)
    #document.add_picture('.\picture\\1.jpg')  # 添加图片
    document.save('.\\report\\'+output_filename)
    line+=1





# document.add_heading('Document Title',
#                      0)  # 这里是给文档添加一个标题，0表示 样式为title，1则为忽略，其他则是Heading{level},具体可以去<a href="https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html" target="_blank">官网</a>查;
#
# p = document.add_paragraph('A plain paragraph having some ')  # 这里是添加一个段落
# p.add_run('bold').bold = True  # 这里是在这个段落p里文字some后面添加bold字符
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
#
# document.add_heading('Heading, level 1', level=1)  # 这里是添加标题1
# document.add_paragraph('Intense quote', style='IntenseQuote')  # 这里是添加段落，style后面则是样式
#
# document.add_paragraph(
#     'first item in unordered list', style='ListBullet'  # 添加段落，样式为unordered list类型
# )
# document.add_paragraph(
#     'first item in ordered list', style='ListNumber'  # 添加段落，样式为ordered list数字类型
# )
#
# document.add_picture('monty-truth.png', width=Inches(1.25))  # 添加图片
#
# table = document.add_table(rows=1, cols=3)  # 添加一个表格，每行三列
# hdr_cells = table.rows[0].cells  # 表格第一行的所含有的所有列数
# hdr_cells[0].text = 'Qty'  # 第一行的第一列,给这行里面添加文字
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for item in recordset:
#     row_cells = table.add_row().cells  # 这是在这个表格第一行 (称作最后一行更好) 下面再添加新的一行
#     row_cells[0].text = str(item.qty)
#     row_cells[1].text = str(item.id)
#     row_cells[2].text = item.desc
#
# document.add_page_break()  # 添加分页符
# #
# # document.save('demo.docx')  # 保存这个文档
